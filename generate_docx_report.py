import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("GV PORTFOLIO ENGINE & CYBER VAULT\nEXECUTIVE SYSTEM REPORT")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(3, 7, 18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Comprehensive Technical Architecture, Security Specifications & User Navigation Guide — Version 5.9.6\nDomain: ganeshvarma.in")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 110, 125)

    doc.add_paragraph() # Spacer

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Core Stack", level=1)
    p1 = doc.add_paragraph(
        "GV Portfolio Engine (v5.9.6) is a modern, high-performance web platform built on the Next.js 14 App Router (React 18 / Node.js). "
        "It presents an interactive tri-mode portfolio ecosystem for Creative Direction, Data Analytics, and Software Engineering, "
        "coupled with a 6-layer defense-in-depth cyber vault security architecture."
    )

    # Core Tech Stack Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component Layer"
    hdr_cells[1].text = "Technology Used"
    hdr_cells[2].text = "Implementation & Features"
    for cell in hdr_cells:
        set_cell_background(cell, "030712")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    stack_data = [
        ("Framework", "Next.js 14 App Router (React 18)", "Server-Side Rendering (SSR) & Dynamic API routes"),
        ("Styling & Design System", "Vanilla CSS & HSL Tokens", "Custom glassmorphism tokens, backdrop blur (24px), dynamic themes"),
        ("Canvas / WebGL Engine", "HTML5 Canvas Vector Physics", "Constellation particle graph, cursor spark trails, vector connections"),
        ("Sound Synthesis", "Browser Web Audio API", "Real-time dual-oscillator sound synthesis (sine, sawtooth, square waves)"),
        ("Authentication & Vault", "NextAuth.js & JWT", "Session encryption, 6-digit keypad PIN, pattern handshake tokens")
    ]

    for layer, tech, impl in stack_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = impl

    doc.add_paragraph() # Spacer

    # Section 2: Security Architecture
    doc.add_heading("2. Multi-Layer Admin Cyber Vault Security Architecture", level=1)
    doc.add_paragraph(
        "To protect administrative surfaces from unauthorized access, automated web crawlers, and brute-force attacks, "
        "the application enforces a 6-Layer Defense-in-Depth Cyber Vault Architecture:"
    )

    sec_layers = [
        ("Layer 1: Secret URL Key Router Guard (/admin?key=134214)", 
         "Direct address bar requests to /admin without the secret URL parameter (?key=YOUR_KEY) are intercepted and served an authentic 404 Page Not Found error template."),
        
        ("Layer 2: Cryptographic Session Pattern Handshake (3-Min TTL)", 
         "Completing the secret star constellation pattern generates a timestamped verification token in sessionStorage (expiresAt: Date.now() + 180000). The token automatically expires after 3 minutes."),
        
        ("Layer 3: 4-Star Multi-Stroke Harry Potter Spell Rune Engine", 
         "4 secret floating star nodes in bounded outer safe zones require an exact 6-stroke sequence (1 -> 3 -> 4 -> 2 -> 1 -> 4). This yields 4^6 = 4,096 exact sequence combinations."),
        
        ("Layer 4: 6-Digit Holographic Cyber Security Keypad Gatekeeper", 
         "Accessing /admin/login presents a 6-digit holographic cyber keypad requiring a secret PIN (Default: 134214) before credential inputs are revealed."),
        
        ("Layer 5: Indian IT Act Cyber Strobe Lockdown Warning & Emergency Siren", 
         "3 incorrect PIN or spell sequence attempts trigger a full-screen cyber strobe lockdown citing Section 66, 66B, & 66F of the Indian IT Act, 2000 & BNS/IPC with a 90s countdown and a dual-oscillator emergency siren."),
        
        ("Layer 6: Authentic 404 Page Not Found Shield", 
         "Renders an authentic 404 Page Not Found layout to unauthorized address bar requests, completely hiding the existence of the Admin Panel.")
    ]

    for title, desc in sec_layers:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}: ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0, 102, 204)
        p.add_run(desc)

    doc.add_paragraph() # Spacer

    # Section 3: Interactive Tri-Mode Portfolio Ecosystem
    doc.add_heading("3. Interactive Tri-Mode Portfolio Ecosystem", level=1)
    doc.add_paragraph(
        "The public application provides 3 specialized, interactive modes tailored for different executive personas:"
    )

    modes = [
        ("🎬 Editor Mode", "Includes a 4K LUT Color Grading Studio with real-time video color presets (Cinema, Cyberpunk, OLED Dark, Vintage Gold) embedded inside a video monitor deck with zero-scroll reel previewing."),
        ("📊 Data Analyst Mode", "Includes a real-time Data Science & AI Ticker Sandbox displaying live metrics, customizable data feeds, and interactive analytics charts."),
        ("💻 Software Developer Mode", "Includes a tabbed executive interactive suite featuring a Cyber AI Assistant terminal, a 2D Gravity Physics Matrix, a System Flow Architecture Blueprint diagram, and a Live UI Transformer.")
    ]

    for m_title, m_desc in modes:
        p = doc.add_paragraph()
        r_m = p.add_run(f"{m_title}: ")
        r_m.bold = True
        p.add_run(m_desc)

    doc.add_paragraph() # Spacer

    # Section 4: Admin Customization & Navigation Guide
    doc.add_heading("4. Admin Customization & Site Navigation Guide", level=1)
    doc.add_paragraph(
        "How to Access the Admin Vault:\n"
        "• Method A (Star Constellation): On the landing page, tap Star 1 -> Star 3 -> Star 4 -> Star 2 -> Star 1 -> Star 4. Click 'ENTER ADMIN PANEL' and enter PIN 134214.\n"
        "• Method B (Direct Secret URL): Navigate to https://www.ganeshvarma.in/admin?key=134214 and enter PIN 134214."
    )

    doc.add_paragraph(
        "Admin Control Surface (/admin/features):\n"
        "Administrators can configure the Secret URL Key, 6-Digit Keypad PIN, 4-Star Spell Rune Sequence, Strike Attempts Limit, and Lockdown Timer duration directly inside /admin/features."
    )

    # Save document
    output_path = "/Users/ganeshvarma/Desktop/gv-next-2/GV_Portfolio_System_Security_Report.docx"
    doc.save(output_path)
    print(f"Successfully created docx report at: {output_path}")

if __name__ == "__main__":
    create_report()
